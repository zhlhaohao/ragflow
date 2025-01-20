#  Licensed under the Apache License, Version 2.0 (the "License");
#  you may not use this file except in compliance with the License.
#  You may obtain a copy of the License at
#
#      http://www.apache.org/licenses/LICENSE-2.0
#
#  Unless required by applicable law or agreed to in writing, software
#  distributed under the License is distributed on an "AS IS" BASIS,
#  WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
#  See the License for the specific language governing permissions and
#  limitations under the License.
#

from docx import Document
import re
import pandas as pd
from collections import Counter
from rag.nlp import rag_tokenizer
from io import BytesIO


class RAGFlowDocxParser:
    """
    调用docx库解析docx
    """

    def __extract_table_content(self, tb):
        """
        提取表格内容并转换为DataFrame格式。

        参数:
        tb (Table): docx库中的表格对象。

        返回:
        list: 处理后的表格内容列表。
        """
        df = []
        for row in tb.rows:
            df.append([c.text for c in row.cells])
        return self.__compose_table_content(pd.DataFrame(df))

    def __compose_table_content(self, df):
        """
        根据表格内容生成结构化的文本。

        参数:
        df (DataFrame): 包含表格数据的DataFrame对象。

        返回:
        list: 结构化的文本列表。
        """
        def blockType(b):
            patt = [
                ("^(20|19)[0-9]{2}[年/-][0-9]{1,2}[月/-][0-9]{1,2}日*$", "Dt"),
                (r"^(20|19)[0-9]{2}年$", "Dt"),
                (r"^(20|19)[0-9]{2}[年/-][0-9]{1,2}月*$", "Dt"),
                ("^[0-9]{1,2}[月/-][0-9]{1,2}日*$", "Dt"),
                (r"^第*[一二三四1-4]季度$", "Dt"),
                (r"^(20|19)[0-9]{2}年*[一二三四1-4]季度$", "Dt"),
                (r"^(20|19)[0-9]{2}[ABCDE]$", "DT"),
                ("^[0-9.,+%/ -]+$", "Nu"),
                (r"^[0-9A-Z/\._~-]+$", "Ca"),
                (r"^[A-Z]*[a-z' -]+$", "En"),
                (r"^[0-9.,+-]+[0-9A-Za-z/$￥%<>（）()' -]+$", "NE"),
                (r"^.{1}$", "Sg")
            ]
            for p, n in patt:
                if re.search(p, b):
                    return n
            tks = [t for t in rag_tokenizer.tokenize(b).split() if len(t) > 1]
            if len(tks) > 3:
                if len(tks) < 12:
                    return "Tx"
                else:
                    return "Lx"

            if len(tks) == 1 and rag_tokenizer.tag(tks[0]) == "nr":
                return "Nr"

            return "Ot"

        if len(df) < 2:
            return []
        max_type = Counter([blockType(str(df.iloc[i, j])) for i in range(
            1, len(df)) for j in range(len(df.iloc[i, :]))])
        max_type = max(max_type.items(), key=lambda x: x[1])[0]

        colnm = len(df.iloc[0, :])
        hdrows = [0]  # header is not necessarily appear in the first line
        if max_type == "Nu":
            for r in range(1, len(df)):
                tys = Counter([blockType(str(df.iloc[r, j]))
                              for j in range(len(df.iloc[r, :]))])
                tys = max(tys.items(), key=lambda x: x[1])[0]
                if tys != max_type:
                    hdrows.append(r)

        lines = []
        for i in range(1, len(df)):
            if i in hdrows:
                continue
            hr = [r - i for r in hdrows]
            hr = [r for r in hr if r < 0]
            t = len(hr) - 1
            while t > 0:
                if hr[t] - hr[t - 1] > 1:
                    hr = hr[t:]
                    break
                t -= 1
            headers = []
            for j in range(len(df.iloc[i, :])):
                t = []
                for h in hr:
                    x = str(df.iloc[i + h, j]).strip()
                    if x in t:
                        continue
                    t.append(x)
                t = ",".join(t)
                if t:
                    t += ": "
                headers.append(t)
            cells = []
            for j in range(len(df.iloc[i, :])):
                if not str(df.iloc[i, j]):
                    continue
                cells.append(headers[j] + str(df.iloc[i, j]))
            lines.append(";".join(cells))

        if colnm > 3:
            return lines
        return ["\n".join(lines)]

    def __call__(self, fnm, from_page=0, to_page=100000000):
        """
        解析docx文件并提取内容。

        参数:
        fnm (str or bytes): 文件路径或文件字节流。
        from_page (int): 开始解析的页码，默认为0。
        to_page (int): 结束解析的页码，默认为一个很大的数。

        返回:
        tuple: 包含段落内容和表格内容的元组。
        """
        # 根据输入的文件路径或字节流创建Document对象
        self.doc = Document(fnm) if isinstance(fnm, str) else Document(BytesIO(fnm))
        pn = 0  # 当前解析的页码
        secs = []  # 存储解析的内容

        # 遍历文档中的每个段落
        for p in self.doc.paragraphs:
            # 如果当前页码超过结束页码，则停止解析
            if pn > to_page:
                break

            runs_within_single_paragraph = []  # 存储当前段落中的run对象
            # 遍历段落中的每个run对象
            for run in p.runs:
                # 如果当前页码超过结束页码，则停止解析
                if pn > to_page:
                    break
                # 如果当前页码在指定范围内且段落文本不为空，则将run文本添加到列表中
                if from_page <= pn < to_page and p.text.strip():
                    runs_within_single_paragraph.append(run.text)  # 先添加run.text

                # 检查run对象中是否包含分页符，如果包含则页码加1
                if 'lastRenderedPageBreak' in run._element.xml:
                    pn += 1

            # 将当前段落中的所有run文本连接成一个字符串，并添加到解析内容列表中
            secs.append(("".join(runs_within_single_paragraph), p.style.name if hasattr(p.style, 'name') else ''))  # 然后将run.text连接成段落

        # 提取文档中的所有表格内容
        tbls = [self.__extract_table_content(tb) for tb in self.doc.tables]
        return secs, tbls
