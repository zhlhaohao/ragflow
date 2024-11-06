#  Licensed under the Apache License, Version 2.0 (the "License");
#  you may not use this file except in compliance with the License.
#  You may obtain a copy of the License at
#
#      http://www.apache.org/licenses/LICENSE-2.0
#
#  Unless required by applicable law or agreed to in writing, software
#  distributed under the License is distributed on an "AS IS" BASIS,
#  WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
#  See the License for the specific language governing permissions and
#  limitations under the License.
#
from tika import parser
from io import BytesIO
from docx import Document
from timeit import default_timer as timer
import re
from deepdoc.parser.pdf_parser import PlainParser
from rag.nlp import rag_tokenizer, naive_merge, tokenize_table, tokenize_chunks, find_codec, concat_img, \
    naive_merge_docx, tokenize_chunks_docx
from deepdoc.parser import PdfParser, ExcelParser, DocxParser, HtmlParser, JsonParser, MarkdownParser, TxtParser
from rag.settings import cron_logger
from rag.utils import num_tokens_from_string
from PIL import Image
from functools import reduce
from markdown import markdown
from docx.image.exceptions import UnrecognizedImageError, UnexpectedEndOfFileError, InvalidImageStreamError


class Docx(DocxParser):
    """
    解析docx,使用docx库
    """
    def __init__(self):
        pass

    def get_picture(self, document, paragraph):
        """
        从段落中提取图片。它首先检查段落中是否有图片，如果有，则尝试通过图片的嵌入ID获取图片的数据流。如果成功，它会将图片数据转换成 PIL.Image 对象并返回。如果遇到错误（如图片格式不被识别、读取过程中遇到意外结束等），则会打印一条消息并返回 None。
        """
        img = paragraph._element.xpath('.//pic:pic')
        if not img:
            return None
        img = img[0]
        embed = img.xpath('.//a:blip/@r:embed')[0]
        related_part = document.part.related_parts[embed]
        try:
            image_blob = related_part.image.blob
        except UnrecognizedImageError:
            print("Unrecognized image format. Skipping image.")
            return None
        except UnexpectedEndOfFileError:
            print("EOF was unexpectedly encountered while reading an image stream. Skipping image.")
            return None
        except InvalidImageStreamError:
            print("The recognized image stream appears to be corrupted. Skipping image.")
            return None
        try:
            image = Image.open(BytesIO(image_blob)).convert('RGB')
            return image
        except Exception as e:
            return None

    def __clean(self, line):
        """
        清理段落中的文本。它将全角空格 \u3000 替换为普通的空格，并去除字符串两端的空白字符
        """
        line = re.sub(r"\u3000", " ", line).strip()
        return line

    def __call__(self, filename, binary=None, from_page=0, to_page=100000):
        """
        实现了主要的解析逻辑。它接受文件名或二进制数据作为输入，解析文档的内容，并返回两个列表：
        1. 一个是包含每个段落的文本和段落内部的所有图片拼接在一起的大图片的元组的列表，
        2. 一个是包含表格HTML表示的元组的列表。

        文档中的每个段落都会被处理，如果段落中有文本，则会被清理并存储在 lines 列表中，同时也会尝试从段落中提取图片。
        如果段落样式为 Caption，则会特殊处理，尝试将最近的图片与之关联。
        当遇到页面断点时（由段落中的运行元素中的 lastRenderedPageBreak 或 w:br 标签指示），页码计数器 pn 会增加。
        最后，所有图片列表被合并成单个图片对象（如果有的话），并且所有的表格被转换成HTML格式。
        """


        # 根据是否提供二进制数据来加载文档
        self.doc = Document(
            filename) if not binary else Document(BytesIO(binary))
        
        # 初始化页码计数器和结果列表
        pn = 0
        lines = []
        last_image = None
        
        # 遍历文档中的每个段落
        for p in self.doc.paragraphs:
            if pn > to_page:
                # 如果当前页码超过指定的结束页码，停止处理
                break
            
            if from_page <= pn < to_page:
                # 段落中有文本内容
                if p.text.strip():
                    if p.style and p.style.name == 'Caption':
                        # 特殊处理 Caption 样式的段落
                        former_image = None
                        if lines and lines[-1][1] and lines[-1][2] != 'Caption':
                            # 获取前一个段落的最后一个图片
                            former_image = lines[-1][1].pop()
                        elif last_image:
                            # 使用上一个未处理的图片
                            former_image = last_image
                            last_image = None
                        # 添加 Caption 段落    
                        lines.append((self.__clean(p.text), [former_image], p.style.name))
                    else:
                        # 处理普通段落
                        # 尝试从段落中提取图片
                        current_image = self.get_picture(self.doc, p)
                        image_list = [current_image]
                        if last_image:
                            # 如果有未处理的图片，添加到当前图片列表前面
                            image_list.insert(0, last_image)
                            last_image = None
                        # 添加普通段落
                        lines.append((self.__clean(p.text), image_list, p.style.name if p.style else ""))
                else:
                    # 段落中无文本，但可能有图片
                    if current_image := self.get_picture(self.doc, p):
                        if lines:
                            # 将图片添加到前一个段落的图片列表中
                            lines[-1][1].append(current_image)
                        else:
                            # 如果没有前一个段落，保存图片以备后续处理
                            last_image = current_image
            # 处理段落中的页面断点
            for run in p.runs:
                if 'lastRenderedPageBreak' in run._element.xml:
                    pn += 1
                    continue
                if 'w:br' in run._element.xml and 'type="page"' in run._element.xml:
                    pn += 1

        # 合并每个段落的图片列表,也就是将该段落的多张图片合并为一张大图片，如果没有图片，则为None
        new_line = [(line[0], reduce(concat_img, line[1]) if line[1] else None) for line in lines]

        # 处理文档中的表格
        tbls = []
        for tb in self.doc.tables:
            html = "<table>"
            for r in tb.rows:
                # 开始新行
                html += "<tr>"
                i = 0
                while i < len(r.cells):
                    span = 1
                    c = r.cells[i]
                    for j in range(i + 1, len(r.cells)):
                        if c.text == r.cells[j].text:
                            # 计算单元格合并的列数
                            span += 1
                            i = j
                    i += 1
                    # 添加单元格
                    html += f"<td>{c.text}</td>" if span == 1 else f"<td colspan='{span}'>{c.text}</td>"
                # 结束行    
                html += "</tr>"
            # 结束表格    
            html += "</table>"
            # 将表格的 HTML 表示添加到结果列表中
            tbls.append(((None, html), ""))

        # 返回处理后的文本和图片列表，以及表格列表
        return new_line, tbls

class Pdf(PdfParser):
    def __call__(self, filename, binary=None, from_page=0,
                 to_page=100000, zoomin=3, callback=None):
        start = timer()
        callback(msg="OCR is running...")
        self.__images__(
            filename if not binary else binary,
            zoomin,
            from_page,
            to_page,
            callback
        )
        callback(msg="OCR finished")
        cron_logger.info("OCR({}~{}): {}".format(from_page, to_page, timer() - start))

        start = timer()
        self._layouts_rec(zoomin)
        callback(0.63, "Layout analysis finished.")
        self._table_transformer_job(zoomin)
        callback(0.65, "Table analysis finished.")
        self._text_merge()
        callback(0.67, "Text merging finished")
        tbls = self._extract_table_figure(True, zoomin, True, True)
        # self._naive_vertical_merge()
        self._concat_downward()
        # self._filter_forpages()

        cron_logger.info("layouts: {}".format(timer() - start))
        return [(b["text"], self._line_tag(b, zoomin))
                for b in self.boxes], tbls


class Markdown(MarkdownParser):
    def __call__(self, filename, binary=None):
        if binary:
            encoding = find_codec(binary)
            txt = binary.decode(encoding, errors="ignore")
        else:
            with open(filename, "r") as f:
                txt = f.read()
        remainder, tables = self.extract_tables_and_remainder(f'{txt}\n')
        sections = []
        tbls = []
        for sec in remainder.split("\n"):
            if num_tokens_from_string(sec) > 10 * self.chunk_token_num:
                sections.append((sec[:int(len(sec) / 2)], ""))
                sections.append((sec[int(len(sec) / 2):], ""))
            else:
                if sections and sections[-1][0].strip().find("#") == 0:
                    sec_, _ = sections.pop(-1)
                    sections.append((sec_+"\n"+sec, ""))
                else:
                    sections.append((sec, ""))

        for table in tables:
            tbls.append(((None, markdown(table, extensions=['markdown.extensions.tables'])), ""))
        return sections, tbls


def chunk(filename, binary=None, from_page=0, to_page=100000,
          lang="Chinese", callback=None, **kwargs):
    """
        Supported file formats are docx, pdf, excel, txt.
        This method apply the naive ways to chunk files.
        Successive text will be sliced into pieces using 'delimiter'.
        Next, these successive pieces are merge into chunks whose token number is no more than 'Max token number'.
    """

    eng = lang.lower() == "english"  # is_english(cks)
    parser_config = kwargs.get(
        "parser_config", {
            "chunk_token_num": 128, "delimiter": "\n!?。；！？", "layout_recognize": True})
    doc = {
        "docnm_kwd": filename,
        "title_tks": rag_tokenizer.tokenize(re.sub(r"\.[a-zA-Z]+$", "", filename))
    }
    doc["title_sm_tks"] = rag_tokenizer.fine_grained_tokenize(doc["title_tks"])
    res = []
    pdf_parser = None

    # 如果文件类型是docx
    if re.search(r"\.docx$", filename, re.IGNORECASE):
        callback(0.1, "Start to parse.")

        # 解析docx，返回解析后的段落（段落文本+段落内图片）列表和表格列表     
        sections, tbls = Docx()(filename, binary)
        # 使用 tokenize_table 方法处理表格。
        res = tokenize_table(tbls, doc, eng)  # just for table

        callback(0.8, "Finish parsing.")
        st = timer()

        # 将段落合并为chunks
        chunks, images = naive_merge_docx(
            sections, int(parser_config.get(
                "chunk_token_num", 128)), parser_config.get(
                "delimiter", "\n!?。；！？"))

        # section_only=True 直接返回段落,但是不会发生
        if kwargs.get("section_only", False):
            return chunks

        # 将chunks tokenize后，附加到res并返回
        res.extend(tokenize_chunks_docx(chunks, doc, eng, images))
        cron_logger.info("naive_merge({}): {}".format(filename, timer() - st))
        return res

    # 如果是pdf文件
    elif re.search(r"\.pdf$", filename, re.IGNORECASE):

        # 如果 layout_recognize=True，那么调用OCR，否则直接从pdf提取文本(PlainParser)
        pdf_parser = Pdf(
        ) if parser_config.get("layout_recognize", True) else PlainParser()
        sections, tbls = pdf_parser(filename if not binary else binary,
                                    from_page=from_page, to_page=to_page, callback=callback)
        # 使用 tokenize_table 方法处理表格。
        res = tokenize_table(tbls, doc, eng)

    elif re.search(r"\.xlsx?$", filename, re.IGNORECASE):
        callback(0.1, "Start to parse.")
        excel_parser = ExcelParser()
        if parser_config.get("html4excel"):
            sections = [(_, "") for _ in excel_parser.html(binary, 12) if _]
        else:
            sections = [(_, "") for _ in excel_parser(binary) if _]

    elif re.search(r"\.(txt|py|js|java|c|cpp|h|php|go|ts|sh|cs|kt|sql)$", filename, re.IGNORECASE):
        callback(0.1, "Start to parse.")
        sections = TxtParser()(filename, binary,
                               parser_config.get("chunk_token_num", 128),
                               parser_config.get("delimiter", "\n!?;。；！？"))
        callback(0.8, "Finish parsing.")

    elif re.search(r"\.(md|markdown)$", filename, re.IGNORECASE):
        callback(0.1, "Start to parse.")
        sections, tbls = Markdown(int(parser_config.get("chunk_token_num", 128)))(filename, binary)
        res = tokenize_table(tbls, doc, eng)
        callback(0.8, "Finish parsing.")

    elif re.search(r"\.(htm|html)$", filename, re.IGNORECASE):
        callback(0.1, "Start to parse.")
        sections = HtmlParser()(filename, binary)
        sections = [(_, "") for _ in sections if _]
        callback(0.8, "Finish parsing.")

    elif re.search(r"\.json$", filename, re.IGNORECASE):
        callback(0.1, "Start to parse.")
        sections = JsonParser(int(parser_config.get("chunk_token_num", 128)))(binary)
        sections = [(_, "") for _ in sections if _]
        callback(0.8, "Finish parsing.")

    elif re.search(r"\.doc$", filename, re.IGNORECASE):
        callback(0.1, "Start to parse.")
        binary = BytesIO(binary)
        doc_parsed = parser.from_buffer(binary)
        sections = doc_parsed['content'].split('\n')
        sections = [(_, "") for _ in sections if _]
        callback(0.8, "Finish parsing.")

    else:
        raise NotImplementedError(
            "file type not supported yet(pdf, xlsx, doc, docx, txt supported)")

    st = timer()
    chunks = naive_merge(
        sections, int(parser_config.get(
            "chunk_token_num", 128)), parser_config.get(
            "delimiter", "\n!?。；！？"))
    if kwargs.get("section_only", False):
        return chunks

    res.extend(tokenize_chunks(chunks, doc, eng, pdf_parser))
    cron_logger.info("naive_merge({}): {}".format(filename, timer() - st))
    return res


if __name__ == "__main__":
    import sys


    def dummy(prog=None, msg=""):
        pass


    chunk(sys.argv[1], from_page=0, to_page=10, callback=dummy)
